#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from model import model
#打开文档
document = Document()
#加入不同等级的标题
exam=model.Exam_model.getByPK(1)
document.add_heading(u'\t\t'+exam['ex_name'],0)

information = model.Information_model()
information_data = information.query('SELECT * FROM information WHERE \
            student_st_id=%s and exam_ex_id=%s' % (2014112207, 2))
information = model.Information_model(**information_data[0])
exam_question = model.Exam_question_model.getByArgs(information_in_id=information.in_id)
choice_question = []
judge_question = []
# filla是读程序写结果
filla_question = []
fillb_question = []
coding_question = []
for item in exam_question:
    if item.eq_qt_type == 'choice':
        question_data = model.Question_model.getByPK(item.qt_id)
        choice_data = model.Choice_model.getByPK(item.qt_id)
        item = dict(item, **choice_data)
        item = dict(item, **question_data)
        choice_question.append(item)
    elif item.eq_qt_type == 'judge':
        question_data = model.Question_model.getByPK(item.qt_id)
        judge_data = model.Judge_model.getByPK(item.qt_id)
        item = dict(item, **judge_data)
        item = dict(item, **question_data)
        judge_question.append(item)
    elif item.eq_qt_type == 'filla':
        question_data = model.Question_model.getByPK(item.qt_id)
        filla_data = model.Filla_model.getByPK(item.qt_id)
        item = dict(item, **filla_data)
        item = dict(item, **question_data)
        filla_question.append(item)
    elif item.eq_qt_type == 'coding':
        question_data = model.Question_model.getByPK(item.qt_id)
        coding_data = model.Coding_model.getByPK(item.qt_id)
        item = dict(item, **coding_data)
        item = dict(item, **question_data)
        coding_question.append(item)
fillb_qt = model.Exam_question_model.query('select distinct(qt_id),\
                                                        eq_pre_score from exam_question where information_in_id = %s \
                                                        and eq_qt_type = %s' % (information.in_id, "'" + 'fillb' + "'"))
for item in fillb_qt:
    question_data = model.Question_model.getByPK(item['qt_id'])
    fillb_data = model.Fillb_model.getByPK(item['qt_id'])

    eq_id_data = model.Exam_question_model.query('select eq_id,eq_answer ,eq_get_score from \
                                                                    exam_question where qt_id = %s and information_in_id = %s' \
                                                     % (item['qt_id'], information.in_id))
    eq_id_data = [model.Exam_question_model(**items) for items in eq_id_data]
    i = 1
    fillb_coding = fillb_data.fb_pre_coding.split('&&&')
    print len(fillb_coding)
    for k in eq_id_data:
        j = 2 * i - 1
        fillb_coding[j] = u'空 '+ str(i)
        i+=1
    fillb_data.fb_pre_coding = ''.join(fillb_coding)

    item = dict(item, **fillb_data)
    item = dict(item, **question_data)
    fillb = []
    fillb.append(item)
    fillb.append(eq_id_data)
        # print fillb
    fillb_question.append(fillb)

student = model.Student_model.getByPK(2014112207)
#添加文本
paragraph = document.add_paragraph()
paragraph.add_run(u'学号:\t')
paragraph.add_run(u'%s'%student['st_id'])
paragraph.add_run(u'\t姓名:    ')
paragraph.add_run(student['st_name'])
paragraph.add_run(u'\t班级:    ')
paragraph.add_run(student['st_specialty'])
paragraph.add_run(u'\t得分:    ')
if information['in_score']<0:
    # information['in_score']=u'未出分'
    paragraph.add_run(u'未出分')
else:
    paragraph.add_run(str(information['in_score']))
# #设置字号
# run = paragraph.add_run(u'设置字号、')
# run.font.size = Pt(24)
#
# #设置字体
# run = paragraph.add_run('Set Font,')
# run.font.name = 'Consolas'

#设置中文字体
# run = paragraph.add_run('学号: '+student['st_id'])
# run.font.name=u'宋体'
# r = run._element
# r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# #增加无序列表
# document.add_paragraph(
#     u'无序列表元素1', style='List Bullet'
# )
# document.add_paragraph(
#     u'无序列表元素2', style='List Bullet'
# )
document.add_heading(u'选择题',1)
for item in choice_question:
    # print item
    if item['eq_get_score']<0:
        # item['eq_get_score'] = "未出分"
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
            style='List Number')
    else:
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
            style='List Number')

    document.add_paragraph(
    u'A. '+item['cc_a'], style='List Bullet'
    )
    document.add_paragraph(
    u'B. ' +item['cc_b'], style='List Bullet'
    )
    document.add_paragraph(
    u'C. ' +item['cc_c'], style='List Bullet'
    )
    document.add_paragraph(
    u'D. ' +item['cc_d'], style='List Bullet'
    )

document.add_heading(u'判断题',1)
for item in judge_question:
    # print item
    if item['eq_get_score']<0:
        # item['eq_get_score'] = "未出分"
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
            style='List Number')
    else:
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
            style='List Number')

document.add_heading(u'读程序写结果',1)
for item in filla_question:
    # print item
    if item['eq_get_score']<0:
        # item['eq_get_score'] = "未出分"
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
            style='List Number')
    else:
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
            style='List Number')





document.add_heading(u'程序填空',1)
for item in fillb_question:
    print item
    answer = ''
    i=1
    for eq in item[1]:
        if eq['eq_get_score'] < 0:
            # eq['eq_get_score'] = u"未出分"
            answer += u'空'+str(i)+ '\t' + eq.eq_answer+u'\t 得分: '+u"未出分" +'\n'
        else:
            answer += u'空' + str(i) + '\t' + eq.eq_answer + u'\t 得分: ' + str(eq['eq_get_score']) + '\n'
        i+=1
    document.add_paragraph(
        item[0]['qt_stem']+u'\n' + item[0]['fb_pre_coding'] +u'\n' +answer, style='List Number'
    )




document.add_heading(u'编程题',1)
for item in coding_question:
    # print item
    if item['eq_get_score']<0:
        # item['eq_get_score'] = "未出分"
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
            style='List Number')
    else:
        document.add_paragraph(
            item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
            style='List Number')

# #增加有序列表
# document.add_paragraph(
#     u'有序列表元素1', style='List Number'
# )



#增加分页
# document.add_page_break()

#保存文件
document.save(u'测试.docx')